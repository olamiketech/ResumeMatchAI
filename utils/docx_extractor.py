from docx import Document
import streamlit as st

def extract_text_from_docx(docx_path):
    """
    Extract text from a DOCX file
    
    Args:
        docx_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from the DOCX
    """
    try:
        doc = Document(docx_path)
        text = ""
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
        
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def validate_docx(docx_path):
    """
    Validate if the DOCX file is readable
    
    Args:
        docx_path (str): Path to the DOCX file
        
    Returns:
        bool: True if DOCX is valid, False otherwise
    """
    try:
        doc = Document(docx_path)
        # Try to access the first paragraph
        if len(doc.paragraphs) > 0:
            doc.paragraphs[0].text
            return True
        return False
    except:
        return False